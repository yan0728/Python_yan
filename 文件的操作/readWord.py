#!/usr/bin/env python 
"""
@author:闫学雷
@project:PythonYan
@file: readWord.py
@time:2020/12/29 0029
"""


import docx

file = docx.Document("111.docx")

# docx 仅支持.docx的文件，直接通过doc后缀改成docx也不好用
# 输出每一段内容
#输出每一段的内容
# for para in file.paragraphs:
    # print(para.text)

#输出段落编号及段落内容
for i in range(len(file.paragraphs)):
    print("第"+str(i)+"段的内容是："+file.paragraphs[i].text)